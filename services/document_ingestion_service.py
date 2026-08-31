import json
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


class DocumentIngestionService:
    """
    Document Ingestion Service

    Loads external insurance documents,
    performs chunking,
    and converts them into LangChain
    Document objects.
    """

    def load_json_documents(
        self,
        json_file_path: str
    ):
        """
        Load structured JSON insurance documents.
        """

        with open(
            json_file_path,
            "r",
            encoding="utf-8"
        ) as f:

            records = json.load(f)

        documents = []

        for record in records:

            content = (
                f"{record.get('title', '')}\n\n"
                f"{record.get('content', '')}"
            )

            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "id":
                            record.get("id"),

                        "category":
                            record.get("category"),

                        "title":
                            record.get("title"),

                        "source":
                            Path(
                                json_file_path
                            ).name
                    }
                )
            )

        return documents

    def chunk_documents(
        self,
        documents
    ):
        """
        Convert large documents
        into smaller retrieval chunks.
        """

        splitter = (
            RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=100
            )
        )

        return (
            splitter.split_documents(
                documents
            )
        )

    ### TXT Loader

    def load_txt_document(
        self,
        txt_file_path: str
    ):
        """
        Load plain text documents.
        """

        with open(
            txt_file_path,
            "r",
            encoding="utf-8"
        ) as f:

            content = f.read()

        return [
            Document(
                page_content=content,
                metadata={
                    "source":
                        Path(
                            txt_file_path
                        ).name,

                    "type":
                        "TXT"
                }
            )
        ]

    ### PDF Loader

    def load_pdf_document(
        self,
        pdf_file_path: str
    ):
        """
        Load PDF documents.
        """

        reader = PdfReader(
            pdf_file_path
        )

        text = ""

        for page in reader.pages:

            extracted = (
                page.extract_text()
            )

            if extracted:

                text += (
                    extracted
                    + "\n"
                )

        return [
            Document(
                page_content=text,
                metadata={
                    "source":
                        Path(
                            pdf_file_path
                        ).name,

                    "type":
                        "PDF"
                }
            )
        ]

    ### DOCX Loader

    def load_docx_document(
        self,
        docx_file_path: str
    ):
        """
        Load Microsoft Word documents.
        """

        document = DocxDocument(
            docx_file_path
        )

        text = "\n".join(
            paragraph.text
            for paragraph
            in document.paragraphs
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source":
                        Path(
                            docx_file_path
                        ).name,

                    "type":
                        "DOCX"
                }
            )
        ]

    def load_directory_documents(
        self,
        directory_path: str
    ):
        """
        Automatically ingest all supported
        files from a directory.

        Supported:
        - JSON
        - TXT
        - PDF
        - DOCX
        """

        folder = Path(
            directory_path
        )

        documents = []

        if not folder.exists():

            return documents

        for file_path in folder.rglob("*"):

            # Ignore folders
            if not file_path.is_file():

                continue

            suffix = (
                file_path.suffix
                .lower()
            )

            try:

                # -------------------------
                # JSON
                # -------------------------

                if suffix == ".json":

                    documents.extend(
                        self.load_json_documents(
                            str(file_path)
                        )
                    )

                # -------------------------
                # TXT
                # -------------------------

                elif suffix == ".txt":

                    documents.extend(
                        self.load_txt_document(
                            str(file_path)
                        )
                    )

                # -------------------------
                # PDF
                # -------------------------

                elif suffix == ".pdf":

                    documents.extend(
                        self.load_pdf_document(
                            str(file_path)
                        )
                    )

                # -------------------------
                # DOCX
                # -------------------------

                elif suffix == ".docx":

                    documents.extend(
                        self.load_docx_document(
                            str(file_path)
                        )
                    )

            except Exception as e:

                print(
                    f"[INGESTION WARNING] "
                    f"{file_path.name}: "
                    f"{e}"
                )

        return documents


document_ingestion_service = (
    DocumentIngestionService()
)